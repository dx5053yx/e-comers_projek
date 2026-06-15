from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "buku-panduan-sipandu.docx"
ASSETS = ROOT / "tmp" / "docx-assets"

GREEN = "149653"
DARK_GREEN = "0B2B20"
SOFT_GREEN = "EAF7F0"
PALE_GREEN = "F5FBF7"
GOLD = "D7A928"
PALE_GOLD = "FFF8E1"
GRAY = "5E6E67"
LIGHT_GRAY = "F3F6F4"
BORDER = "C9D9D0"
WHITE = "FFFFFF"
RED = "B42318"
PALE_RED = "FEF3F2"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name="Aptos", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def add_title(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    set_keep_with_next(paragraph)
    return paragraph


def add_body(document, text, bold_prefix=None):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        set_run_font(run)


def add_steps(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        set_run_font(run)


def add_callout(document, title, text, tone="green"):
    palettes = {
        "green": (SOFT_GREEN, GREEN),
        "gold": (PALE_GOLD, GOLD),
        "red": (PALE_RED, RED),
        "gray": (LIGHT_GRAY, GRAY),
    }
    fill, accent = palettes[tone]
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, accent, "10")
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    p = cell.paragraphs[0]
    title_run = p.add_run(title)
    set_run_font(title_run, bold=True, color=accent)
    p.add_run("\n")
    text_run = p.add_run(text)
    set_run_font(text_run, color=DARK_GREEN if tone != "red" else RED)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_two_column_cards(document, cards):
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index in range(0, len(cards), 2):
        cells = table.add_row().cells
        for column in range(2):
            cell = cells[column]
            cell.width = Cm(8.3)
            set_cell_margins(cell, top=170, start=180, bottom=170, end=180)
            if index + column >= len(cards):
                set_cell_border(cell, WHITE, "0")
                continue
            title, text = cards[index + column]
            set_cell_shading(cell, PALE_GREEN)
            set_cell_border(cell)
            title_p = cell.paragraphs[0]
            title_run = title_p.add_run(title)
            set_run_font(title_run, size=11, bold=True, color=GREEN)
            body_p = cell.add_paragraph()
            body_p.paragraph_format.space_after = Pt(0)
            body_run = body_p.add_run(text)
            set_run_font(body_run, size=9.5, color=DARK_GREEN)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_feature_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header_cells = table.rows[0].cells
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = header_cells[index]
        if widths:
            cell.width = Cm(widths[index])
        set_cell_shading(cell, DARK_GREEN)
        set_cell_border(cell, DARK_GREEN)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        set_run_font(run, size=9.5, bold=True, color=WHITE)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            if widths:
                cell.width = Cm(widths[index])
            set_cell_shading(cell, WHITE if row_index % 2 == 0 else PALE_GREEN)
            set_cell_border(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            set_run_font(run, size=9.2, color=DARK_GREEN)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_product_gallery(document):
    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    products = [
        (
            "Seblak Original",
            "Rp10.000",
            "Kerupuk, telur, sayuran, dan kuah pedas.",
            ASSETS / "seblak-original.png",
        ),
        (
            "Seblak Seafood",
            "Rp15.000",
            "Seafood, bakso ikan, telur, dan kuah pedas.",
            ASSETS / "seblak-seafood.png",
        ),
        (
            "Es Teh",
            "Rp4.000",
            "Es teh manis segar sebagai pendamping.",
            ASSETS / "es-teh.png",
        ),
    ]
    for index, (name, price, description, image_path) in enumerate(products):
        cell = table.cell(0, index)
        cell.width = Cm(5.5)
        set_cell_shading(cell, WHITE)
        set_cell_border(cell)
        set_cell_margins(cell, top=100, start=100, bottom=150, end=100)
        image_p = cell.paragraphs[0]
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_p.add_run().add_picture(str(image_path), width=Cm(5.0))
        title_p = cell.add_paragraph()
        title_run = title_p.add_run(name)
        set_run_font(title_run, size=10.5, bold=True, color=DARK_GREEN)
        price_p = cell.add_paragraph()
        price_run = price_p.add_run(price)
        set_run_font(price_run, size=10, bold=True, color=GREEN)
        desc_p = cell.add_paragraph()
        desc_p.paragraph_format.space_after = Pt(0)
        desc_run = desc_p.add_run(description)
        set_run_font(desc_run, size=8.5, color=GRAY)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_status_flow(document):
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    statuses = [
        ("1", "Menunggu Bayar"),
        ("2", "Diproses"),
        ("3", "Dikemas"),
        ("4", "Dikirim"),
        ("5", "Selesai"),
    ]
    for index, (number, label) in enumerate(statuses):
        cell = table.cell(0, index)
        cell.width = Cm(3.2)
        set_cell_shading(cell, SOFT_GREEN if index < 4 else PALE_GOLD)
        set_cell_border(cell, GREEN if index < 4 else GOLD)
        set_cell_margins(cell, top=120, start=80, bottom=120, end=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        number_run = p.add_run(number)
        set_run_font(number_run, size=13, bold=True, color=GREEN)
        p.add_run("\n")
        label_run = p.add_run(label)
        set_run_font(label_run, size=8.5, bold=True, color=DARK_GREEN)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    document.core_properties.title = "Buku Panduan Penggunaan siPandu"
    document.core_properties.subject = "Panduan penggunaan platform e-commerce siPandu"
    document.core_properties.author = "siPandu"
    document.core_properties.keywords = "siPandu, e-commerce, UMKM, WhatsApp, katalog, pesanan"

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        style.font.size = Pt(10.2)
        style.font.color.rgb = RGBColor.from_string(DARK_GREEN)

    heading_1 = styles["Heading 1"]
    heading_1.font.name = "Aptos Display"
    heading_1._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    heading_1.font.size = Pt(20)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    heading_1.paragraph_format.space_before = Pt(8)
    heading_1.paragraph_format.space_after = Pt(10)

    heading_2 = styles["Heading 2"]
    heading_2.font.name = "Aptos Display"
    heading_2._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    heading_2.font.size = Pt(14)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor.from_string(GREEN)
    heading_2.paragraph_format.space_before = Pt(10)
    heading_2.paragraph_format.space_after = Pt(5)

    heading_3 = styles["Heading 3"]
    heading_3.font.name = "Aptos"
    heading_3._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    heading_3.font.size = Pt(11.5)
    heading_3.font.bold = True
    heading_3.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    heading_3.paragraph_format.space_before = Pt(7)
    heading_3.paragraph_format.space_after = Pt(3)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text = footer_p.add_run("Buku Panduan siPandu  |  ")
    set_run_font(footer_text, size=8, color=GRAY)
    add_page_number(footer_p)

    # Cover
    cover_logo = document.add_paragraph()
    cover_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_logo.add_run().add_picture(str(ASSETS / "logo.png"), width=Cm(2.6))
    cover_logo.paragraph_format.space_after = Pt(12)

    cover_title = document.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = cover_title.add_run("BUKU PANDUAN\nPENGGUNAAN siPandu")
    set_run_font(title_run, name="Aptos Display", size=28, bold=True, color=DARK_GREEN)

    cover_subtitle = document.add_paragraph()
    cover_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = cover_subtitle.add_run(
        "Platform E-Commerce dan Layanan Penjualan WhatsApp untuk UMKM"
    )
    set_run_font(subtitle_run, size=13, color=GREEN)
    cover_subtitle.paragraph_format.space_after = Pt(18)

    hero_p = document.add_paragraph()
    hero_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hero_p.add_run().add_picture(str(ROOT / "public" / "hero-sipandu.png"), width=Cm(16.5))
    hero_p.paragraph_format.space_after = Pt(14)

    cover_business = document.add_table(rows=1, cols=1)
    cover_business.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_cell = cover_business.cell(0, 0)
    set_cell_shading(cover_cell, SOFT_GREEN)
    set_cell_border(cover_cell, GREEN, "10")
    set_cell_margins(cover_cell, top=180, start=220, bottom=180, end=220)
    cover_business_p = cover_cell.paragraphs[0]
    cover_business_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    business_run = cover_business_p.add_run("Contoh Penggunaan: Warung Seblak Ibu Ani")
    set_run_font(business_run, size=12, bold=True, color=DARK_GREEN)

    cover_date = document.add_paragraph()
    cover_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_date.paragraph_format.space_before = Pt(16)
    date_run = cover_date.add_run("Edisi Juni 2026")
    set_run_font(date_run, size=10, color=GRAY)
    document.add_page_break()

    # Intro and contents
    add_title(document, "Petunjuk Penggunaan Buku", 1)
    add_body(
        document,
        "Buku ini menjelaskan cara menggunakan siPandu dari sudut pandang pengguna "
        "e-commerce. Panduan dibagi untuk pelanggan, pemilik UMKM, dan tamu yang "
        "ingin melihat dashboard saat presentasi.",
    )
    add_callout(
        document,
        "Tujuan utama",
        "Membantu pengguna memahami proses mencari produk, membuat pesanan, membayar, "
        "memantau pengiriman, mengelola toko, menjalankan promo, serta melayani "
        "pelanggan melalui WhatsApp.",
        "green",
    )

    add_title(document, "Daftar Isi", 1)
    contents = [
        ("1", "Mengenal siPandu"),
        ("2", "Jenis Pengguna dan Hak Akses"),
        ("3", "Panduan Pelanggan"),
        ("4", "Panduan Pemilik UMKM"),
        ("5", "Mengelola Profil, Katalog, dan Stok"),
        ("6", "Mengelola Pesanan, Pembayaran, dan Pengiriman"),
        ("7", "Promo, Pelanggan, Ulasan, dan Analitik"),
        ("8", "Layanan WhatsApp dan AI"),
        ("9", "Akun Tamu untuk Presentasi"),
        ("10", "Penggunaan di Ponsel dan Pilihan Tema"),
        ("11", "Skenario Demonstrasi Warung Seblak Ibu Ani"),
        ("12", "Pertanyaan Umum dan Penyelesaian Masalah"),
        ("13", "Checklist Operasional"),
    ]
    add_feature_table(document, ["Bab", "Isi Panduan"], contents, [2, 14.5])
    document.add_page_break()

    # 1
    add_title(document, "1. Mengenal siPandu", 1)
    add_body(
        document,
        "siPandu adalah platform e-commerce untuk UMKM yang menggabungkan katalog "
        "produk, pemesanan, pembayaran, pengiriman, promosi, data pelanggan, ulasan, "
        "analitik penjualan, dan layanan WhatsApp dalam satu alur kerja.",
    )
    add_two_column_cards(
        document,
        [
            ("Katalog digital", "Produk, foto, harga, varian, dan ketersediaan dapat dilihat pelanggan."),
            ("Pesanan terpusat", "Order dari katalog dan WhatsApp tampil pada dashboard toko."),
            ("Pembayaran QRIS", "Pelanggan menerima instruksi pembayaran dan dapat mengirim bukti bayar."),
            ("Pengiriman", "Pemilik toko dapat mencatat kurir, nomor resi, dan status pengiriman."),
            ("Promosi", "Diskon atau promo beli X gratis Y dapat ditawarkan kepada pelanggan."),
            ("Layanan pelanggan", "WhatsApp AI membantu menjawab pertanyaan dan mengarahkan proses pemesanan."),
        ],
    )
    add_title(document, "Alur E-Commerce siPandu", 2)
    add_feature_table(
        document,
        ["Tahap", "Kegiatan Pelanggan", "Kegiatan UMKM"],
        [
            ("1. Temukan toko", "Membuka direktori UMKM atau WhatsApp.", "Menampilkan profil dan katalog."),
            ("2. Pilih produk", "Melihat produk, harga, stok, dan promo.", "Menjaga data produk tetap akurat."),
            ("3. Buat pesanan", "Mengisi produk, jumlah, nama, WhatsApp, dan alamat.", "Menerima order pada dashboard."),
            ("4. Bayar", "Scan QRIS atau transfer dan kirim bukti.", "Memeriksa serta memverifikasi pembayaran."),
            ("5. Diproses", "Memantau status order.", "Menyiapkan dan mengemas pesanan."),
            ("6. Dikirim", "Melihat kurir dan nomor resi.", "Memperbarui informasi pengiriman."),
            ("7. Beri ulasan", "Memberikan rating dan komentar.", "Melihat ulasan sebagai bahan evaluasi."),
        ],
        [3.1, 6.6, 6.6],
    )
    document.add_page_break()

    # 2
    add_title(document, "2. Jenis Pengguna dan Hak Akses", 1)
    add_feature_table(
        document,
        ["Pengguna", "Kebutuhan", "Akses Utama"],
        [
            (
                "Pelanggan",
                "Mencari produk dan melakukan pembelian.",
                "Direktori UMKM, katalog, checkout, pembayaran, tracking, dan ulasan.",
            ),
            (
                "Pemilik/Pengelola UMKM",
                "Menjalankan operasional toko.",
                "Produk, stok, order, pembayaran, pengiriman, promo, pelanggan, WhatsApp, dan analitik.",
            ),
            (
                "Tamu Presentasi",
                "Melihat kemampuan sistem tanpa mengubah data.",
                "Seluruh halaman dashboard dalam mode hanya-baca.",
            ),
        ],
        [3.5, 5.5, 7.3],
    )
    add_callout(
        document,
        "Pelanggan tidak perlu membuat akun",
        "Pelanggan dapat langsung melihat katalog dan membuat pesanan. Login hanya "
        "diperlukan oleh pemilik UMKM atau tamu yang ingin melihat dashboard.",
        "gold",
    )
    add_title(document, "Tanda Mode Tamu", 2)
    add_body(
        document,
        "Saat masuk sebagai tamu, dashboard menampilkan pemberitahuan 'Mode tamu hanya-baca'. "
        "Tamu dapat membuka menu dan melihat data, tetapi tombol perubahan data tidak tersedia "
        "atau dinonaktifkan.",
    )
    document.add_page_break()

    # 3
    add_title(document, "3. Panduan Pelanggan", 1)
    add_title(document, "3.1 Menemukan UMKM", 2)
    add_steps(
        document,
        [
            "Buka halaman utama siPandu.",
            "Pilih menu UMKM untuk melihat usaha yang telah terdaftar.",
            "Gunakan informasi nama toko, kategori, alamat, dan kontak WhatsApp untuk memilih toko.",
            "Buka katalog toko yang diinginkan.",
        ],
    )
    add_title(document, "3.2 Melihat Katalog", 2)
    add_body(
        document,
        "Pada katalog, pelanggan dapat melihat foto produk, nama, deskripsi, harga, dan "
        "status ketersediaan. Contoh katalog Warung Seblak Ibu Ani:",
    )
    add_product_gallery(document)
    add_title(document, "3.3 Membuat Pesanan", 2)
    add_steps(
        document,
        [
            "Pilih produk pada bagian Checkout.",
            "Masukkan jumlah produk yang dipesan.",
            "Isi nama pelanggan.",
            "Isi nomor WhatsApp aktif agar informasi pesanan dapat diterima.",
            "Isi alamat pengiriman atau lokasi pengambilan.",
            "Tambahkan catatan, misalnya tingkat pedas atau waktu pengambilan.",
            "Periksa kembali pilihan lalu tekan tombol untuk membuat pesanan.",
            "Simpan kode order yang ditampilkan.",
        ],
    )
    add_callout(
        document,
        "Periksa nomor WhatsApp",
        "Gunakan nomor yang aktif dan ditulis dengan benar. Nomor tersebut membantu toko "
        "menghubungi pelanggan jika ada perubahan atau konfirmasi pesanan.",
        "gold",
    )
    add_title(document, "3.4 Membayar Pesanan", 2)
    add_steps(
        document,
        [
            "Buka halaman pembayaran melalui link atau kode order.",
            "Periksa total pembayaran.",
            "Scan QRIS toko atau ikuti instruksi transfer.",
            "Lakukan pembayaran sesuai nominal.",
            "Kirim atau upload bukti pembayaran.",
            "Tunggu pengelola toko memverifikasi pembayaran.",
        ],
    )
    add_title(document, "3.5 Melacak Pesanan", 2)
    add_body(
        document,
        "Halaman tracking menampilkan status pembayaran, proses pesanan, kurir, dan nomor "
        "resi. Gunakan kode order untuk membuka kembali informasi tersebut.",
    )
    add_status_flow(document)
    add_title(document, "3.6 Memberikan Ulasan", 2)
    add_steps(
        document,
        [
            "Pastikan pesanan telah selesai.",
            "Buka link ulasan dari WhatsApp atau halaman order.",
            "Pilih rating 1 sampai 5.",
            "Tambahkan komentar mengenai rasa, pelayanan, atau pengiriman.",
            "Kirim ulasan.",
        ],
    )
    document.add_page_break()

    # 4
    add_title(document, "4. Panduan Pemilik UMKM", 1)
    add_title(document, "4.1 Daftar dan Masuk", 2)
    add_steps(
        document,
        [
            "Buka halaman Daftar UMKM.",
            "Isi nama pengelola, email, password, nama toko, dan nomor WhatsApp.",
            "Selesaikan pendaftaran.",
            "Buka halaman Login.",
            "Masukkan email dan password pengelola.",
            "Pilih Masuk sebagai pengelola.",
        ],
    )
    add_title(document, "4.2 Memahami Dashboard", 2)
    add_two_column_cards(
        document,
        [
            ("Total penjualan", "Nilai penjualan dari pesanan yang telah dibayar."),
            ("Total pesanan", "Jumlah seluruh order yang tercatat."),
            ("Menunggu bayar", "Pesanan yang masih memerlukan pembayaran atau verifikasi."),
            ("Diproses", "Pesanan yang sedang disiapkan hingga dikemas."),
            ("Stok menipis", "Produk yang telah menyentuh batas stok minimum."),
            ("Rata-rata ulasan", "Nilai kepuasan pelanggan dari rating yang masuk."),
        ],
    )
    add_title(document, "4.3 Navigasi", 2)
    add_feature_table(
        document,
        ["Kelompok", "Menu", "Kegunaan"],
        [
            ("Utama", "Dashboard, Analitik", "Melihat ringkasan dan perkembangan penjualan."),
            ("Operasional", "Produk, Stok, Pesanan, Pembayaran, Pengiriman", "Menjalankan aktivitas penjualan harian."),
            ("Relasi", "Pelanggan, WhatsApp, Ulasan, Promo", "Menjaga hubungan dan komunikasi pelanggan."),
            ("Sistem", "Pengaturan", "Mengelola profil dan informasi toko."),
        ],
        [3.0, 6.5, 7.0],
    )
    document.add_page_break()

    # 5
    add_title(document, "5. Mengelola Profil, Katalog, dan Stok", 1)
    add_title(document, "5.1 Mengatur Profil Toko", 2)
    add_steps(
        document,
        [
            "Buka menu Pengaturan.",
            "Isi nama bisnis dan kategori usaha.",
            "Isi nomor WhatsApp toko.",
            "Isi alamat serta deskripsi singkat.",
            "Isi instruksi pembayaran.",
            "Upload QRIS toko.",
            "Tekan Simpan setting.",
        ],
    )
    add_callout(
        document,
        "Profil toko adalah etalase digital",
        "Gunakan nama, deskripsi, foto, alamat, dan nomor WhatsApp yang jelas agar pelanggan "
        "percaya dan mudah menghubungi toko.",
        "green",
    )
    add_title(document, "5.2 Menambah Produk", 2)
    add_steps(
        document,
        [
            "Buka menu Produk.",
            "Tekan Tambah produk.",
            "Isi nama produk, kode produk, harga, dan deskripsi.",
            "Upload foto produk yang jelas.",
            "Isi nama varian jika tersedia.",
            "Masukkan jumlah stok.",
            "Atur batas peringatan stok menipis.",
            "Pastikan produk berstatus aktif.",
            "Tekan Simpan produk.",
        ],
    )
    add_feature_table(
        document,
        ["Bagian Produk", "Contoh", "Saran"],
        [
            ("Nama", "Seblak Seafood", "Gunakan nama yang mudah dipahami."),
            ("Harga", "15.000", "Masukkan harga jual akhir."),
            ("Deskripsi", "Seblak dengan topping seafood dan bakso ikan.", "Jelaskan isi dan keunggulan produk."),
            ("Foto", "Foto makanan yang terang", "Gunakan foto asli dan tidak buram."),
            ("Stok", "20", "Perbarui saat persediaan berubah."),
            ("Batas stok", "5", "Sesuaikan dengan waktu belanja bahan."),
        ],
        [3.3, 5.2, 8.0],
    )
    add_title(document, "5.3 Menonaktifkan Produk", 2)
    add_body(
        document,
        "Jika produk sedang tidak dijual, gunakan aksi Hapus dari katalog atau nonaktifkan. "
        "Produk tidak lagi tampil kepada pelanggan, tetapi riwayat pesanan tetap tersimpan.",
    )
    add_title(document, "5.4 Memantau Stok", 2)
    add_bullets(
        document,
        [
            "Status In stock berarti persediaan masih di atas batas minimum.",
            "Status Low stock berarti persediaan perlu segera ditambah.",
            "Stok akan berkurang ketika pembayaran pesanan telah diverifikasi.",
            "Periksa stok sebelum mengaktifkan promo dalam jumlah besar.",
        ],
    )
    document.add_page_break()

    # 6
    add_title(document, "6. Mengelola Pesanan, Pembayaran, dan Pengiriman", 1)
    add_title(document, "6.1 Membuka Pesanan", 2)
    add_steps(
        document,
        [
            "Buka menu Pesanan.",
            "Pilih kode order yang ingin diperiksa.",
            "Periksa nama pelanggan, WhatsApp, alamat, item, jumlah, total, dan catatan.",
            "Gunakan panel antrean untuk berpindah ke order lain tanpa kembali ke daftar.",
        ],
    )
    add_title(document, "6.2 Mengubah Status Pesanan", 2)
    add_status_flow(document)
    add_feature_table(
        document,
        ["Status", "Arti", "Tindakan UMKM"],
        [
            ("Menunggu pembayaran", "Pembayaran belum diterima atau belum diperiksa.", "Tunggu bukti atau hubungi pelanggan."),
            ("Diproses", "Pembayaran telah diterima.", "Mulai menyiapkan produk."),
            ("Packing", "Pesanan sedang dikemas.", "Pastikan item dan jumlah sesuai."),
            ("Dikirim", "Pesanan telah diserahkan ke kurir.", "Isi kurir dan nomor resi."),
            ("Selesai", "Pesanan telah diterima.", "Minta pelanggan memberikan ulasan."),
            ("Dibatalkan", "Pesanan tidak dilanjutkan.", "Catat alasan pembatalan."),
        ],
        [3.8, 6.1, 6.6],
    )
    add_title(document, "6.3 Memverifikasi Pembayaran", 2)
    add_steps(
        document,
        [
            "Buka bukti pembayaran.",
            "Cocokkan nama penerima, tujuan pembayaran, tanggal, dan nominal.",
            "Pilih Verifikasi bayar jika bukti sesuai.",
            "Pilih Tolak pembayaran jika bukti salah atau tidak jelas.",
            "Hubungi pelanggan jika diperlukan.",
        ],
    )
    add_callout(
        document,
        "Jangan verifikasi berdasarkan tulisan saja",
        "Periksa gambar bukti pembayaran dengan teliti. Status PAID menandakan pembayaran "
        "telah diterima dan dapat memengaruhi stok.",
        "red",
    )
    add_title(document, "6.4 Mengatur Pengiriman", 2)
    add_steps(
        document,
        [
            "Buka detail pesanan.",
            "Isi nama kurir atau metode pengantaran.",
            "Isi nomor resi jika tersedia.",
            "Pilih status pengiriman.",
            "Tekan Simpan pengiriman.",
        ],
    )
    document.add_page_break()

    # 7
    add_title(document, "7. Promo, Pelanggan, Ulasan, dan Analitik", 1)
    add_title(document, "7.1 Membuat Promo", 2)
    add_steps(
        document,
        [
            "Buka menu Promo.",
            "Pilih Voucher diskon atau Beli X gratis Y.",
            "Isi nama dan kode promo.",
            "Masukkan nilai diskon atau jumlah beli dan gratis.",
            "Atur minimal belanja atau minimal item.",
            "Atur kuota penggunaan dan periode bila diperlukan.",
            "Aktifkan promo lalu tekan Simpan promo.",
        ],
    )
    add_feature_table(
        document,
        ["Jenis", "Contoh", "Cocok Digunakan Untuk"],
        [
            ("Diskon persen", "HEMAT10: diskon 10%", "Mendorong transaksi pada periode tertentu."),
            ("Diskon nominal", "POTONG5K: potongan Rp5.000", "Memberi manfaat yang mudah dipahami."),
            ("Beli X gratis Y", "Beli 3 gratis 1", "Meningkatkan jumlah item per pesanan."),
        ],
        [4.0, 5.0, 7.5],
    )
    add_title(document, "7.2 Data Pelanggan", 2)
    add_body(
        document,
        "Menu Pelanggan membantu UMKM mengenali pelanggan baru, pelanggan yang kembali "
        "berbelanja, dan pelanggan loyal. Gunakan informasi ini untuk pelayanan dan promosi "
        "yang lebih sesuai.",
    )
    add_title(document, "7.3 Ulasan", 2)
    add_bullets(
        document,
        [
            "Baca rating dan komentar pelanggan secara berkala.",
            "Gunakan kritik untuk memperbaiki produk atau pelayanan.",
            "Pertahankan produk yang mendapat respons baik.",
            "Jangan mengubah ulasan pelanggan menjadi informasi yang menyesatkan.",
        ],
    )
    add_title(document, "7.4 Analitik Penjualan", 2)
    add_body(
        document,
        "Analitik membantu pemilik UMKM melihat pola penjualan, produk terlaris, total order, "
        "penjualan per periode, serta kondisi stok. Informasi ini dapat digunakan untuk "
        "menentukan jumlah persediaan dan jenis promo berikutnya.",
    )
    document.add_page_break()

    # 8
    add_title(document, "8. Layanan WhatsApp dan AI", 1)
    add_title(document, "8.1 Menghubungkan WhatsApp", 2)
    add_steps(
        document,
        [
            "Buka menu WhatsApp pada dashboard.",
            "Tunggu kode QR tersedia.",
            "Buka WhatsApp bisnis pada ponsel.",
            "Pilih Perangkat tertaut.",
            "Pilih Tautkan perangkat.",
            "Scan kode QR yang tampil pada dashboard.",
            "Tunggu status berubah menjadi Terhubung.",
        ],
    )
    add_title(document, "8.2 Hal yang Dapat Dibantu AI", 2)
    add_two_column_cards(
        document,
        [
            ("Menjawab pertanyaan", "Menjelaskan produk, harga, stok, dan cara pesan."),
            ("Menawarkan produk", "Memberikan rekomendasi yang relevan secara wajar."),
            ("Menawarkan promo", "Menyebut promo aktif ketika pesanan memenuhi syarat."),
            ("Mencatat order", "Mengarahkan percakapan hingga pesanan tercatat."),
            ("Mengirim pembayaran", "Menyampaikan total dan QRIS setelah order dibuat."),
            ("Meminta rating", "Mengajak pelanggan memberi ulasan setelah pesanan selesai."),
        ],
    )
    add_title(document, "8.3 Mengatur Gaya Balasan", 2)
    add_body(
        document,
        "Pemilik toko dapat mengisi gaya balasan agar AI berbicara sesuai karakter bisnis. "
        "Contoh: ramah, santai, menggunakan sapaan 'Kak', menawarkan menu populer, dan "
        "mengarahkan pelanggan untuk membuat pesanan tanpa memaksa.",
    )
    add_callout(
        document,
        "Informasi tetap harus akurat",
        "AI sebaiknya hanya menyampaikan produk, harga, stok, promo, pembayaran, dan status "
        "pesanan yang tersedia pada sistem.",
        "gold",
    )
    add_title(document, "8.4 Contoh Percakapan", 2)
    chat_table = document.add_table(rows=0, cols=1)
    chat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    messages = [
        ("Pelanggan", "Halo, ada menu apa saja?", LIGHT_GRAY, DARK_GREEN),
        (
            "siPandu",
            "Halo Kak. Hari ini tersedia Seblak Original Rp10.000, Seblak Seafood "
            "Rp15.000, dan Es Teh Rp4.000. Kakak mau yang pedas biasa atau seafood?",
            SOFT_GREEN,
            GREEN,
        ),
        ("Pelanggan", "Seblak Seafood 2 dan Es Teh 1.", LIGHT_GRAY, DARK_GREEN),
        (
            "siPandu",
            "Siap Kak. Pesanannya 2 Seblak Seafood dan 1 Es Teh. Aku bantu catat "
            "pesanannya ya. Boleh kirim nama dan alamat?",
            SOFT_GREEN,
            GREEN,
        ),
    ]
    for sender, message, fill, accent in messages:
        cell = chat_table.add_row().cells[0]
        set_cell_shading(cell, fill)
        set_cell_border(cell, accent)
        set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
        p = cell.paragraphs[0]
        sender_run = p.add_run(sender)
        set_run_font(sender_run, size=9, bold=True, color=accent)
        p.add_run("\n")
        message_run = p.add_run(message)
        set_run_font(message_run, size=10, color=DARK_GREEN)
    document.add_page_break()

    # 9
    add_title(document, "9. Akun Tamu untuk Presentasi", 1)
    add_body(
        document,
        "Akun tamu membantu dosen, penguji, atau teman sekelas melihat dashboard "
        "Warung Seblak Ibu Ani tanpa risiko mengubah data operasional.",
    )
    add_steps(
        document,
        [
            "Buka halaman Login.",
            "Pilih Lihat dashboard sebagai tamu.",
            "Tunggu dashboard Warung Seblak Ibu Ani terbuka.",
            "Gunakan menu untuk melihat produk, stok, pesanan, pembayaran, pengiriman, pelanggan, promo, ulasan, WhatsApp, dan analitik.",
            "Pilih Keluar setelah selesai.",
        ],
    )
    add_feature_table(
        document,
        ["Bisa Dilakukan", "Tidak Bisa Dilakukan"],
        [
            ("Melihat ringkasan penjualan", "Menambah atau mengubah produk"),
            ("Membuka daftar order", "Mengubah status order"),
            ("Melihat stok dan promo", "Memverifikasi pembayaran"),
            ("Melihat pelanggan dan ulasan", "Menghubungkan atau memutus WhatsApp"),
            ("Melihat analitik", "Mengubah profil toko atau gaya AI"),
        ],
        [8.2, 8.2],
    )
    add_callout(
        document,
        "Aman untuk demonstrasi",
        "Jika tombol perubahan data tidak tersedia saat memakai akun tamu, itu adalah "
        "perilaku yang benar.",
        "green",
    )
    document.add_page_break()

    # 10
    add_title(document, "10. Penggunaan di Ponsel dan Pilihan Tema", 1)
    add_title(document, "10.1 Navigasi Ponsel", 2)
    add_steps(
        document,
        [
            "Tekan tombol menu di kanan atas dashboard.",
            "Pilih kategori dan halaman yang ingin dibuka.",
            "Tekan area di luar panel atau tombol tutup untuk menutup menu.",
        ],
    )
    add_title(document, "10.2 Mode Gelap dan Terang", 2)
    add_body(
        document,
        "Gunakan tombol Mode gelap atau Mode terang sesuai kenyamanan. Pada desktop, "
        "kontrol tema berada di bagian bawah sidebar. Pada ponsel, kontrol tersedia di "
        "dalam menu dashboard.",
    )
    add_bullets(
        document,
        [
            "Mode terang cocok digunakan di ruangan terang.",
            "Mode gelap membantu mengurangi silau pada ruangan redup.",
            "Pilihan tema akan tersimpan pada perangkat yang digunakan.",
        ],
    )
    document.add_page_break()

    # 11
    add_title(document, "11. Skenario Demonstrasi Warung Seblak Ibu Ani", 1)
    add_body(
        document,
        "Skenario berikut dapat digunakan untuk menunjukkan alur e-commerce siPandu secara "
        "ringkas dan mudah dipahami.",
    )
    add_feature_table(
        document,
        ["Urutan", "Demonstrasi", "Hal yang Dijelaskan"],
        [
            ("1", "Buka halaman utama", "Tujuan siPandu bagi UMKM dan pelanggan."),
            ("2", "Buka direktori UMKM", "Pelanggan dapat menemukan toko lokal."),
            ("3", "Buka katalog Warung Seblak Ibu Ani", "Foto, harga, produk, dan checkout."),
            ("4", "Buat pesanan Seblak Seafood dan Es Teh", "Alur pemesanan pelanggan."),
            ("5", "Tampilkan pembayaran QRIS", "Proses pembayaran dan bukti bayar."),
            ("6", "Masuk sebagai tamu", "Dashboard dapat dilihat tanpa mengubah data."),
            ("7", "Buka order, payment, dan shipment", "Operasional penjualan terpusat."),
            ("8", "Tampilkan promo dan analitik", "Strategi penjualan dan evaluasi bisnis."),
            ("9", "Tampilkan percakapan WhatsApp", "AI membantu pelayanan dan pencatatan order."),
            ("10", "Tampilkan ulasan pelanggan", "Penilaian layanan setelah transaksi selesai."),
        ],
        [2.0, 6.2, 8.3],
    )
    add_callout(
        document,
        "Saran presentasi",
        "Mulai dari sudut pandang pelanggan, lalu beralih ke dashboard pemilik UMKM. "
        "Urutan ini menunjukkan perjalanan transaksi secara utuh.",
        "gold",
    )
    document.add_page_break()

    # 12
    add_title(document, "12. Pertanyaan Umum dan Penyelesaian Masalah", 1)
    faq = [
        (
            "Mengapa foto produk tidak muncul?",
            "Muat ulang halaman. Jika masih kosong, pemilik UMKM perlu membuka menu Produk dan memastikan foto sudah tersimpan.",
        ),
        (
            "Mengapa produk tidak tersedia di katalog?",
            "Periksa apakah produk berstatus aktif dan memiliki stok.",
        ),
        (
            "Mengapa pesanan belum masuk tahap proses?",
            "Pembayaran mungkin belum diverifikasi. Periksa menu Pembayaran atau detail order.",
        ),
        (
            "Mengapa status masih Menunggu Bayar setelah bukti dikirim?",
            "Bukti pembayaran harus diperiksa dan diverifikasi oleh pengelola toko.",
        ),
        (
            "Mengapa QR WhatsApp tidak muncul?",
            "Tunggu beberapa saat lalu gunakan tombol Hubungkan atau Reset QR. Jika tetap tidak muncul, hubungi pengelola sistem.",
        ),
        (
            "Mengapa akun tamu tidak bisa mengedit?",
            "Akun tamu memang hanya untuk melihat dashboard dan tidak memiliki hak mengubah data.",
        ),
        (
            "Bagaimana pelanggan memberikan rating?",
            "Buka link ulasan setelah order selesai, pilih rating 1 sampai 5, tulis komentar, lalu kirim.",
        ),
        (
            "Bagaimana mengganti tampilan gelap atau terang?",
            "Gunakan kontrol tema di sidebar desktop atau di dalam menu ponsel.",
        ),
    ]
    for question, answer in faq:
        add_title(document, question, 3)
        add_body(document, answer)
    document.add_page_break()

    # 13
    add_title(document, "13. Checklist Operasional", 1)
    add_title(document, "Sebelum Toko Mulai Menerima Pesanan", 2)
    add_bullets(
        document,
        [
            "Profil toko, alamat, dan nomor WhatsApp sudah benar.",
            "Foto serta deskripsi produk sudah lengkap.",
            "Harga dan stok sudah diperiksa.",
            "QRIS dan instruksi pembayaran sudah tersedia.",
            "Promo aktif sudah sesuai periode dan syarat.",
            "WhatsApp toko sudah terhubung.",
        ],
    )
    add_title(document, "Setiap Ada Pesanan Baru", 2)
    add_bullets(
        document,
        [
            "Periksa item, jumlah, alamat, dan catatan pelanggan.",
            "Periksa bukti pembayaran sebelum verifikasi.",
            "Perbarui status pesanan sesuai kondisi sebenarnya.",
            "Isi kurir dan nomor resi ketika pesanan dikirim.",
        ],
    )
    add_title(document, "Setelah Pesanan Selesai", 2)
    add_bullets(
        document,
        [
            "Pastikan status pesanan menjadi Selesai.",
            "Arahkan pelanggan untuk memberikan rating.",
            "Baca ulasan yang masuk.",
            "Periksa analitik dan stok untuk rencana penjualan berikutnya.",
        ],
    )
    add_callout(
        document,
        "Pelayanan yang konsisten",
        "Data yang diperbarui dengan disiplin membuat katalog, chatbot, stok, dan dashboard "
        "menyampaikan informasi yang sama kepada pelanggan.",
        "green",
    )

    final_p = document.add_paragraph()
    final_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_p.paragraph_format.space_before = Pt(24)
    final_run = final_p.add_run("siPandu - Membantu UMKM Melayani, Menjual, dan Bertumbuh")
    set_run_font(final_run, size=12, bold=True, color=GREEN)

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output_path = build_document()
    print(output_path)
